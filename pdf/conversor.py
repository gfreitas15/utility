from __future__ import annotations

import os
import time
from datetime import datetime
from pathlib import Path

from PyQt5.QtCore import Qt, QTimer, QFileSystemWatcher
from PyQt5.QtGui import QFont
from PyQt5.QtWidgets import (
    QCheckBox,
    QComboBox,
    QFileDialog,
    QGroupBox,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QListWidget,
    QMessageBox,
    QPushButton,
    QSizePolicy,
    QSplitter,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)

from shared.pdf_deps import (
    PIL_AVAILABLE,
    Image,
    PdfMerger,
    PdfReader,
    PdfWriter,
    openpyxl,
    Document,
    pdf2image,
    canvas,
    A4,
)


class ConversorPDFWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.pasta_monitoramento = ""
        self.pasta_saida = ""  # Não definir pasta padrão, será definida quando necessário
        self.watcher = None
        self.monitorando = False
        self.arquivos_convertidos = set()  # Usar set para melhor performance
        self.timer_monitoramento = None  # Timer de backup
        self.ultima_verificacao = 0  # Timestamp da última verificação
        self.init_ui()
        self.setAcceptDrops(True)

        # Aplicar tema escuro por padrão na inicialização
        self.aplicar_tema(True)

        # Inicializar log se as dependências estiverem disponíveis
        if PIL_AVAILABLE and hasattr(self, "text_log"):
            self.adicionar_log("Conversor de PDF iniciado")

    def init_ui(self):
        if not PIL_AVAILABLE:
            self._mostrar_erro_dependencias()
            return

        layout = QVBoxLayout()
        layout.setSpacing(15)  # Espaçamento entre seções

        # Splitter principal para dividir a interface (handles invisíveis)
        splitter_principal = QSplitter(Qt.Horizontal)
        splitter_principal.setHandleWidth(1)
        splitter_principal.setChildrenCollapsible(False)
        splitter_principal.setStyleSheet("QSplitter::handle { background-color: transparent; }")

        # Painel esquerdo - Conversão de Imagens
        painel_esquerdo = self._criar_painel_conversao()
        splitter_principal.addWidget(painel_esquerdo)

        # Splitter direito para dividir junção e conversões especiais (handles invisíveis)
        splitter_direito = QSplitter(Qt.Vertical)
        splitter_direito.setHandleWidth(1)
        splitter_direito.setChildrenCollapsible(False)
        splitter_direito.setStyleSheet("QSplitter::handle { background-color: transparent; }")

        # Painel superior direito - Conversões Especiais
        painel_conversoes_especiais = self._criar_painel_conversoes_especiais()
        splitter_direito.addWidget(painel_conversoes_especiais)

        # Painel inferior direito - Junção de PDFs e Logs
        painel_juncao_logs = self._criar_painel_juncao_logs()
        splitter_direito.addWidget(painel_juncao_logs)

        splitter_principal.addWidget(splitter_direito)

        # Configurar proporções dos splitters
        splitter_principal.setSizes([400, 400])
        splitter_direito.setSizes([200, 300])
        layout.addWidget(splitter_principal)

        self.setLayout(layout)

    def _mostrar_erro_dependencias(self):
        layout = QVBoxLayout()
        msg = QLabel(
            "❌ Dependências não instaladas!\n\nPara usar o Conversor de PDF, instale as dependências:\n\npip install Pillow reportlab PyPDF2"
        )
        msg.setAlignment(Qt.AlignCenter)
        msg.setFont(QFont("Segoe UI", 12))
        msg.setStyleSheet("color: red; padding: 20px;")
        layout.addWidget(msg)
        self.setLayout(layout)

    def _criar_painel_conversao(self):
        group = QGroupBox("🖼️ Conversão de Imagens para PDF")
        layout = QVBoxLayout()
        layout.setSpacing(10)  # Espaçamento interno

        # Monitoramento de pasta
        monitor_group = QGroupBox("📁 Monitoramento de Pasta")
        monitor_layout = QVBoxLayout()

        # Seleção de pasta de entrada (monitoramento)
        entrada_layout = QHBoxLayout()
        self.btn_pasta_entrada = QPushButton("📂 Selecionar Entrada")
        self.btn_pasta_entrada.clicked.connect(self.selecionar_pasta_entrada)
        self.lbl_pasta_entrada = QLabel("Nenhuma pasta de entrada selecionada")
        self.lbl_pasta_entrada.setWordWrap(True)
        entrada_layout.addWidget(self.btn_pasta_entrada)
        monitor_layout.addLayout(entrada_layout)
        monitor_layout.addWidget(self.lbl_pasta_entrada)

        # Seleção de pasta de saída
        saida_layout = QHBoxLayout()
        self.btn_pasta_saida_monitor = QPushButton("📁 Selecionar Saída")
        self.btn_pasta_saida_monitor.clicked.connect(self.selecionar_pasta_saida_monitor)
        self.lbl_pasta_saida_monitor = QLabel("Nenhuma pasta de saída selecionada")
        self.lbl_pasta_saida_monitor.setWordWrap(True)
        saida_layout.addWidget(self.btn_pasta_saida_monitor)
        monitor_layout.addLayout(saida_layout)
        monitor_layout.addWidget(self.lbl_pasta_saida_monitor)

        # Botões de controle do monitoramento
        botoes_layout = QHBoxLayout()

        self.btn_iniciar_monitor = QPushButton("▶️ Iniciar")
        self.btn_iniciar_monitor.clicked.connect(self.iniciar_monitoramento)
        self.btn_iniciar_monitor.setEnabled(False)

        self.btn_pausar_monitor = QPushButton("⏸️ Pausar")
        self.btn_pausar_monitor.clicked.connect(self.pausar_monitoramento)
        self.btn_pausar_monitor.setEnabled(False)

        self.btn_cancelar_monitor = QPushButton("⏹️ Cancelar")
        self.btn_cancelar_monitor.clicked.connect(self.cancelar_monitoramento)
        self.btn_cancelar_monitor.setEnabled(False)

        self.btn_limpar_monitor = QPushButton("🗑️ Limpar")
        self.btn_limpar_monitor.clicked.connect(self.limpar_monitoramento)
        self.btn_limpar_monitor.setEnabled(True)

        botoes_layout.addWidget(self.btn_iniciar_monitor)
        botoes_layout.addWidget(self.btn_pausar_monitor)
        botoes_layout.addWidget(self.btn_cancelar_monitor)
        botoes_layout.addWidget(self.btn_limpar_monitor)

        monitor_layout.addLayout(botoes_layout)

        # Opção manter original no monitoramento
        self.check_manter_original_monitor = QCheckBox("Manter original")
        self.check_manter_original_monitor.setChecked(True)
        monitor_layout.addWidget(self.check_manter_original_monitor)

        # Status do monitoramento
        self.lbl_status_monitor = QLabel("Status: Parado")
        self.lbl_status_monitor.setStyleSheet("color: red; font-weight: bold;")
        monitor_layout.addWidget(self.lbl_status_monitor)

        monitor_group.setLayout(monitor_layout)
        layout.addWidget(monitor_group)

        # Conversão manual
        manual_group = QGroupBox("📎 Conversão Manual")
        manual_layout = QVBoxLayout()

        # Seleção de arquivos (acima da lista)
        arquivos_layout = QHBoxLayout()
        self.btn_selecionar_imagens = QPushButton("📂 Selecionar Imagens")
        self.btn_selecionar_imagens.clicked.connect(self.selecionar_imagens)
        arquivos_layout.addWidget(self.btn_selecionar_imagens)
        manual_layout.addLayout(arquivos_layout)

        # Lista de arquivos selecionados
        self.lista_imagens = QListWidget()
        # A altura da lista será razoável; botões ficarão abaixo
        self.lista_imagens.setMinimumHeight(180)
        manual_layout.addWidget(self.lista_imagens)

        # Linha de botões pequenos abaixo da lista
        botoes_abaixo_layout = QHBoxLayout()
        self.btn_pasta_saida = QPushButton("📁 Pasta de Saída")
        self.btn_pasta_saida.clicked.connect(self.selecionar_pasta_saida)
        self.btn_pasta_saida.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        botoes_abaixo_layout.addWidget(self.btn_pasta_saida)

        self.btn_limpar_lista = QPushButton("🗑️ Limpar")
        self.btn_limpar_lista.clicked.connect(self.limpar_lista_imagens)
        self.btn_limpar_lista.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        botoes_abaixo_layout.addWidget(self.btn_limpar_lista)

        self.btn_converter_imagens = QPushButton("🔄 Converter Imagens")
        self.btn_converter_imagens.clicked.connect(self.converter_imagens)
        self.btn_converter_imagens.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        botoes_abaixo_layout.addWidget(self.btn_converter_imagens)

        # Opção manter original na conversão manual
        self.check_manter_original_manual = QCheckBox("Manter original")
        self.check_manter_original_manual.setChecked(True)
        botoes_abaixo_layout.addWidget(self.check_manter_original_manual)

        botoes_abaixo_layout.addStretch()
        manual_layout.addLayout(botoes_abaixo_layout)

        # Rótulo de pasta de saída abaixo dos botões
        self.lbl_pasta_saida = QLabel("Nenhuma pasta de saída selecionada")
        self.lbl_pasta_saida.setWordWrap(True)
        manual_layout.addWidget(self.lbl_pasta_saida)

        manual_group.setLayout(manual_layout)
        layout.addWidget(manual_group)

        group.setLayout(layout)
        return group

    def _criar_painel_conversoes_especiais(self):
        group = QGroupBox("🔄 Conversões Especiais")
        layout = QVBoxLayout()
        layout.setSpacing(10)  # Espaçamento interno

        # Seleção de tipo de conversão
        conversao_group = QGroupBox("📋 Escolher Conversão")
        conversao_layout = QVBoxLayout()

        # Dropdowns para seleção
        dropdowns_layout = QHBoxLayout()

        # Dropdown "De"
        de_layout = QVBoxLayout()
        de_layout.addWidget(QLabel("De:"))
        self.cmb_de = QComboBox()
        self.cmb_de.addItems(
            [
                "📊 Excel (.xlsx, .xls)",
                "📄 PDF (.pdf)",
                "📝 Word (.docx)",
                "🖼️ Imagem (.png, .jpg, .jpeg, .gif, .bmp, .tiff, .webp)",
            ]
        )
        self.cmb_de.currentTextChanged.connect(self.atualizar_filtros_arquivo)
        de_layout.addWidget(self.cmb_de)
        dropdowns_layout.addLayout(de_layout)

        # Seta
        seta_label = QLabel("→")
        seta_label.setAlignment(Qt.AlignCenter)
        seta_label.setFont(QFont("Segoe UI", 16, QFont.Bold))
        dropdowns_layout.addWidget(seta_label)

        # Dropdown "Para"
        para_layout = QVBoxLayout()
        para_layout.addWidget(QLabel("Para:"))
        self.cmb_para = QComboBox()
        self.cmb_para.addItems(["📄 PDF (.pdf)", "📝 Word (.docx)", "🖼️ Imagem (.png)", "📊 Excel (.xlsx)"])
        para_layout.addWidget(self.cmb_para)
        dropdowns_layout.addLayout(para_layout)

        conversao_layout.addLayout(dropdowns_layout)

        # Linha de ações compactas logo abaixo (botões pequenos lado a lado)
        acoes_layout = QHBoxLayout()
        self.btn_selecionar_arquivo = QPushButton("📂 Selecionar Arquivo")
        self.btn_selecionar_arquivo.clicked.connect(self.selecionar_arquivo_conversao)
        self.btn_selecionar_arquivo.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        acoes_layout.addWidget(self.btn_selecionar_arquivo)

        # Botão limpar conversão especial ao lado do selecionar
        self.btn_limpar_especial = QPushButton("🗑️ Limpar")
        self.btn_limpar_especial.clicked.connect(self.limpar_conversao_especial)
        self.btn_limpar_especial.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        acoes_layout.addWidget(self.btn_limpar_especial)

        self.btn_converter_especial = QPushButton("🔄 Converter")
        self.btn_converter_especial.clicked.connect(self.converter_arquivo_especial)
        self.btn_converter_especial.setEnabled(False)
        self.btn_converter_especial.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        acoes_layout.addWidget(self.btn_converter_especial)

        self.check_manter_original_especial = QCheckBox("Manter original")
        self.check_manter_original_especial.setChecked(True)
        acoes_layout.addWidget(self.check_manter_original_especial)
        acoes_layout.addStretch()

        conversao_layout.addLayout(acoes_layout)

        # Linha de status do arquivo selecionado
        self.lbl_arquivo_selecionado = QLabel("Nenhum arquivo selecionado")
        self.lbl_arquivo_selecionado.setWordWrap(True)
        conversao_layout.addWidget(self.lbl_arquivo_selecionado)

        conversao_group.setLayout(conversao_layout)
        layout.addWidget(conversao_group)

        group.setLayout(layout)
        return group

    def _criar_painel_juncao_logs(self):
        group = QGroupBox("📚 Junção de PDFs e Logs")
        layout = QVBoxLayout()
        layout.setSpacing(10)  # Espaçamento interno

        # Junção de PDFs
        juncao_group = QGroupBox("🔗 Junção de PDFs")
        juncao_layout = QVBoxLayout()

        # Seleção de PDFs para juntar
        botoes_juncao_layout = QHBoxLayout()
        self.btn_selecionar_pdfs = QPushButton("📂 Selecionar PDFs")
        self.btn_selecionar_pdfs.clicked.connect(self.selecionar_pdfs_para_juntar)
        self.btn_selecionar_pdfs.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        botoes_juncao_layout.addWidget(self.btn_selecionar_pdfs)

        self.btn_limpar_juncao = QPushButton("🗑️ Limpar")
        self.btn_limpar_juncao.clicked.connect(self.limpar_lista_juncao)
        self.btn_limpar_juncao.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        botoes_juncao_layout.addWidget(self.btn_limpar_juncao)

        botoes_juncao_layout.addStretch()
        juncao_layout.addLayout(botoes_juncao_layout)

        # Lista de PDFs selecionados
        self.lista_pdfs = QListWidget()
        self.lista_pdfs.setMaximumHeight(80)
        juncao_layout.addWidget(self.lista_pdfs)

        # Botão de junção
        self.btn_juntar_pdfs = QPushButton("🔗 Juntar PDFs")
        self.btn_juntar_pdfs.clicked.connect(self.juntar_pdfs)
        juncao_layout.addWidget(self.btn_juntar_pdfs)

        juncao_group.setLayout(juncao_layout)
        layout.addWidget(juncao_group)

        # Logs
        logs_group = QGroupBox("📋 Log de Atividades")
        logs_layout = QVBoxLayout()

        # Área de log
        self.text_log = QTextEdit()
        # Aproximadamente 10 linhas de altura
        self.text_log.setMinimumHeight(220)
        self.text_log.setMaximumHeight(260)
        self.text_log.setReadOnly(True)
        logs_layout.addWidget(self.text_log)

        # Botão para limpar log
        self.btn_limpar_log = QPushButton("🗑️ Limpar Log")
        self.btn_limpar_log.clicked.connect(self.limpar_log)
        logs_layout.addWidget(self.btn_limpar_log)

        logs_group.setLayout(logs_layout)
        layout.addWidget(logs_group)

        group.setLayout(layout)
        return group

    # === MÉTODOS DE FUNCIONALIDADE ===

    def adicionar_log(self, mensagem):
        """Adiciona uma mensagem ao log com timestamp"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        log_entry = f"[{timestamp}] {mensagem}"
        self.text_log.append(log_entry)
        # Auto-scroll para o final
        self.text_log.verticalScrollBar().setValue(self.text_log.verticalScrollBar().maximum())

    def limpar_log(self):
        """Limpa o log de atividades"""
        self.text_log.clear()
        self.adicionar_log("Log limpo")

    def selecionar_pasta_entrada(self):
        """Seleciona a pasta de entrada (monitoramento)"""
        pasta = QFileDialog.getExistingDirectory(self, "Selecionar Pasta de Entrada (Monitoramento)")
        if pasta:
            self.pasta_monitoramento = pasta
            self.lbl_pasta_entrada.setText(f"📂 {pasta}")
            self.adicionar_log(f"Pasta de entrada selecionada: {pasta}")

            # Se não há pasta de saída definida, usar a mesma pasta
            if not hasattr(self, "pasta_saida") or not self.pasta_saida:
                self.pasta_saida = pasta
                self.lbl_pasta_saida_monitor.setText(f"📁 {pasta}")
                self.adicionar_log(f"Pasta de saída definida automaticamente como: {pasta}")

            # Habilitar botão Iniciar quando pasta de entrada for selecionada
            self.btn_iniciar_monitor.setEnabled(True)

    def selecionar_pasta_saida_monitor(self):
        """Seleciona a pasta de saída para o monitoramento"""
        pasta = QFileDialog.getExistingDirectory(self, "Selecionar Pasta de Saída (PDFs Convertidos)")
        if pasta:
            self.pasta_saida = pasta
            self.lbl_pasta_saida_monitor.setText(f"📁 {pasta}")
            self.adicionar_log(f"Pasta de saída selecionada: {pasta}")

    def selecionar_pasta_saida(self):
        """Seleciona a pasta de saída dos PDFs"""
        pasta = QFileDialog.getExistingDirectory(self, "Selecionar Pasta de Saída")
        if pasta:
            self.pasta_saida = pasta
            self.lbl_pasta_saida.setText(f"📁 {pasta}")
            self.adicionar_log(f"Pasta de saída personalizada selecionada: {pasta}")
            self.adicionar_log("💡 Dica: Para usar a pasta de monitoramento como saída, cancele e reinicie o monitoramento")

    def toggle_monitoramento(self, state):
        """Método mantido para compatibilidade - não usado mais"""
        pass

    def iniciar_monitoramento(self):
        """Inicia o monitoramento da pasta"""
        if not self.pasta_monitoramento:
            QMessageBox.warning(self, "Erro", "Selecione uma pasta de entrada primeiro!")
            return

        # Se não há pasta de saída selecionada, usar a pasta de entrada
        if not self.pasta_saida:
            self.pasta_saida = self.pasta_monitoramento
            self.lbl_pasta_saida_monitor.setText(f"📁 {self.pasta_saida}")
            self.adicionar_log(f"Pasta de saída definida automaticamente como: {self.pasta_saida}")

        # Limpar watcher anterior se existir
        if self.watcher:
            self.watcher.deleteLater()

        # Criar novo watcher
        self.watcher = QFileSystemWatcher()
        self.watcher.addPath(self.pasta_monitoramento)
        self.watcher.directoryChanged.connect(self.processar_pasta_monitorada)

        # Criar timer de backup (verifica a cada 2 segundos)
        if self.timer_monitoramento:
            self.timer_monitoramento.stop()

        self.timer_monitoramento = QTimer()
        self.timer_monitoramento.timeout.connect(self.verificar_pasta_periodicamente)
        self.timer_monitoramento.start(2000)  # 2 segundos

        self.monitorando = True
        self.ultima_verificacao = time.time()
        self.lbl_status_monitor.setText("Status: Monitorando")
        self.lbl_status_monitor.setStyleSheet("color: green; font-weight: bold;")
        self.adicionar_log(f"Monitoramento iniciado em: {self.pasta_monitoramento}")
        self.adicionar_log("Sistema de monitoramento duplo ativado (watcher + timer)")

        # Atualizar estado dos botões
        self.btn_iniciar_monitor.setEnabled(False)
        self.btn_pausar_monitor.setEnabled(True)
        self.btn_cancelar_monitor.setEnabled(True)

    def pausar_monitoramento(self):
        """Pausa o monitoramento da pasta"""
        if self.watcher:
            self.watcher.deleteLater()
            self.watcher = None

        if self.timer_monitoramento:
            self.timer_monitoramento.stop()

        self.monitorando = False
        self.lbl_status_monitor.setText("Status: Pausado")
        self.lbl_status_monitor.setStyleSheet("color: orange; font-weight: bold;")
        self.adicionar_log("Monitoramento pausado")

        # Atualizar estado dos botões
        self.btn_iniciar_monitor.setEnabled(True)
        self.btn_pausar_monitor.setEnabled(False)
        self.btn_cancelar_monitor.setEnabled(True)

    def cancelar_monitoramento(self):
        """Cancela o monitoramento da pasta"""
        if self.watcher:
            self.watcher.deleteLater()
            self.watcher = None

        if self.timer_monitoramento:
            self.timer_monitoramento.stop()
            self.timer_monitoramento = None

        self.monitorando = False
        self.lbl_status_monitor.setText("Status: Parado")
        self.lbl_status_monitor.setStyleSheet("color: red; font-weight: bold;")
        self.adicionar_log("Monitoramento cancelado")

        # Limpar lista de arquivos convertidos
        self.arquivos_convertidos.clear()
        self.adicionar_log("Lista de arquivos convertidos limpa")

        # Atualizar estado dos botões
        self.btn_iniciar_monitor.setEnabled(True)
        self.btn_pausar_monitor.setEnabled(False)
        self.btn_cancelar_monitor.setEnabled(False)

    def parar_monitoramento(self):
        """Método mantido para compatibilidade - chama cancelar"""
        self.cancelar_monitoramento()

    def limpar_monitoramento(self):
        """Limpa as pastas selecionadas e a lista de arquivos convertidos do monitoramento"""
        if self.monitorando:
            QMessageBox.warning(self, "Aviso", "Pare o monitoramento antes de limpar!")
            return

        # Limpar pastas selecionadas
        self.pasta_monitoramento = ""
        self.pasta_saida = ""
        self.lbl_pasta_entrada.setText("Nenhuma pasta de entrada selecionada")
        self.lbl_pasta_saida_monitor.setText("Nenhuma pasta de saída selecionada")

        # Limpar lista de arquivos convertidos
        self.arquivos_convertidos.clear()

        # Desabilitar botão iniciar
        self.btn_iniciar_monitor.setEnabled(False)

        self.adicionar_log("Monitoramento limpo: pastas e lista de arquivos convertidos")
        QMessageBox.information(
            self,
            "Limpeza",
            "Monitoramento limpo!\n- Pastas de entrada e saída foram limpas\n- Lista de arquivos convertidos foi limpa\n- Agora você pode selecionar novas pastas",
        )

    def processar_pasta_monitorada(self, path):
        """Processa arquivos na pasta monitorada (chamado pelo watcher)"""
        if not self.monitorando:
            return

        self.adicionar_log(f"📁 Mudança detectada na pasta: {path}")
        self._processar_arquivos_na_pasta(path)

    def verificar_pasta_periodicamente(self):
        """Verificação periódica de backup (chamado pelo timer)"""
        if not self.monitorando or not self.pasta_monitoramento:
            return

        self._processar_arquivos_na_pasta(self.pasta_monitoramento)

    def _processar_arquivos_na_pasta(self, path):
        """Processa arquivos de imagem na pasta especificada"""
        try:
            # Lista arquivos de imagem na pasta
            extensoes_imagem = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"}
            arquivos_imagem = []

            if not os.path.exists(path):
                self.adicionar_log(f"⚠️ Pasta não existe: {path}")
                return

            for arquivo in os.listdir(path):
                if Path(arquivo).suffix.lower() in extensoes_imagem:
                    arquivo_path = os.path.join(path, arquivo)
                    if os.path.isfile(arquivo_path):
                        # Verificar se o arquivo já foi convertido (usar caminho completo)
                        if arquivo_path not in self.arquivos_convertidos:
                            # Verificar se o arquivo não está sendo usado por outro processo
                            if self._arquivo_pronto_para_conversao(arquivo_path):
                                arquivos_imagem.append(arquivo_path)

            # Converte cada imagem encontrada
            if arquivos_imagem:
                self.adicionar_log(f"🔄 Encontrados {len(arquivos_imagem)} arquivo(s) para conversão")
                for arquivo_imagem in arquivos_imagem:
                    self.converter_imagem_para_pdf_monitoramento(
                        arquivo_imagem,
                        self.pasta_saida,
                        manter_original=self.check_manter_original_monitor.isChecked(),
                    )
            else:
                # Log apenas ocasionalmente para não poluir
                if time.time() - self.ultima_verificacao > 30:  # A cada 30 segundos
                    self.adicionar_log("👀 Verificando pasta... (nenhum arquivo novo encontrado)")
                    self.ultima_verificacao = time.time()

        except Exception as e:
            self.adicionar_log(f"❌ Erro ao processar pasta: {str(e)}")

    def _arquivo_pronto_para_conversao(self, arquivo_path):
        """Verifica se o arquivo está pronto para conversão (não está sendo usado)"""
        try:
            # Tentar abrir o arquivo para verificar se não está sendo usado
            with open(arquivo_path, "rb") as f:
                f.read(1)  # Ler apenas 1 byte
            return True
        except (PermissionError, OSError):
            # Arquivo está sendo usado por outro processo
            return False

    def selecionar_imagens(self):
        """Seleciona imagens para conversão manual"""
        arquivos, _ = QFileDialog.getOpenFileNames(
            self,
            "Selecionar Imagens",
            "",
            "Imagens (*.png *.jpg *.jpeg *.gif *.bmp *.tiff *.webp);;Todos os arquivos (*)",
        )

        if arquivos:
            for arquivo in arquivos:
                self.lista_imagens.addItem(f"📁 {arquivo}")
            self.adicionar_log(f"{len(arquivos)} imagem(ns) selecionada(s)")

    def limpar_lista_imagens(self):
        """Limpa a lista de imagens selecionadas"""
        self.lista_imagens.clear()
        self.adicionar_log("Lista de imagens limpa")

    def limpar_conversao_especial(self):
        """Limpa seleção e opções da conversão especial"""
        if hasattr(self, "arquivo_especial_selecionado"):
            delattr(self, "arquivo_especial_selecionado")
        self.lbl_arquivo_selecionado.setText("Nenhum arquivo selecionado")
        self.btn_converter_especial.setEnabled(False)
        # mantém as escolhas de De/Para

    def limpar_lista_juncao(self):
        """Limpa lista de PDFs para junção"""
        self.lista_pdfs.clear()
        self.adicionar_log("Lista de PDFs para junção limpa")

    def converter_imagens(self):
        """Converte as imagens selecionadas para PDF"""
        if self.lista_imagens.count() == 0:
            QMessageBox.warning(self, "Erro", "Selecione pelo menos uma imagem!")
            return

        if not self.pasta_saida:
            # Usar Desktop como padrão
            self.pasta_saida = os.path.join(os.path.expanduser("~"), "Desktop")
            self.lbl_pasta_saida.setText(f"📁 {self.pasta_saida}")

        # Converter cada imagem
        for i in range(self.lista_imagens.count()):
            arquivo_imagem = self.lista_imagens.item(i).text()
            # Remover emoji "📁 " do início do caminho
            if arquivo_imagem.startswith("📁 "):
                arquivo_imagem = arquivo_imagem[2:]  # Remove "📁 "
            self.converter_imagem_para_pdf(
                arquivo_imagem, self.pasta_saida, manter_original=self.check_manter_original_manual.isChecked()
            )

        # Limpar lista após conversão
        self.limpar_lista_imagens()

    def converter_imagem_para_pdf(self, caminho_imagem, pasta_saida, manter_original=None):
        """Converte uma imagem para PDF"""
        try:
            if manter_original is None:
                manter_original = True
                if hasattr(self, "check_manter_original_manual"):
                    manter_original = self.check_manter_original_manual.isChecked()
            # Abrir a imagem
            with Image.open(caminho_imagem) as img:
                # Converter para RGB se necessário
                if img.mode != "RGB":
                    img = img.convert("RGB")

                # Criar nome do PDF
                nome_arquivo = Path(caminho_imagem).stem
                caminho_pdf = os.path.join(pasta_saida, f"{nome_arquivo}.pdf")

                # Salvar como PDF
                img.save(caminho_pdf, "PDF", resolution=300.0)

                # Remover arquivo original (respeita flag manter_original)
                if not manter_original:
                    os.remove(caminho_imagem)

                self.adicionar_log(f"✅ Convertido: {Path(caminho_imagem).name} → {Path(caminho_pdf).name}")

        except Exception as e:
            self.adicionar_log(f"❌ Erro ao converter {Path(caminho_imagem).name}: {str(e)}")

    def converter_imagem_para_pdf_monitoramento(self, caminho_imagem, pasta_saida, manter_original=None):
        """Converte uma imagem para PDF com controle de duplicatas para monitoramento"""
        try:
            if manter_original is None:
                manter_original = True
                if hasattr(self, "check_manter_original_monitor"):
                    manter_original = self.check_manter_original_monitor.isChecked()

            # Abrir a imagem
            with Image.open(caminho_imagem) as img:
                # Converter para RGB se necessário
                if img.mode != "RGB":
                    img = img.convert("RGB")

                # Criar nome do PDF com controle de duplicatas
                nome_arquivo = Path(caminho_imagem).stem
                caminho_pdf = self._gerar_nome_arquivo_unico(pasta_saida, nome_arquivo, ".pdf")

                # Salvar como PDF
                img.save(caminho_pdf, "PDF", resolution=300.0)

                # Adicionar à lista de arquivos convertidos (usar caminho completo)
                self.arquivos_convertidos.add(caminho_imagem)

                # Remover arquivo original (respeita flag manter_original)
                if not manter_original:
                    os.remove(caminho_imagem)

                self.adicionar_log(f"✅ Convertido: {Path(caminho_imagem).name} → {Path(caminho_pdf).name}")

        except Exception as e:
            self.adicionar_log(f"❌ Erro ao converter {Path(caminho_imagem).name}: {str(e)}")

    def _gerar_nome_arquivo_unico(self, pasta, nome_base, extensao):
        """Gera um nome de arquivo único, adicionando número se necessário"""
        caminho_base = os.path.join(pasta, f"{nome_base}{extensao}")

        if not os.path.exists(caminho_base):
            return caminho_base

        contador = 2
        while True:
            novo_nome = f"{nome_base}{contador}{extensao}"
            novo_caminho = os.path.join(pasta, novo_nome)
            if not os.path.exists(novo_caminho):
                return novo_caminho
            contador += 1

    def selecionar_pdfs_para_juntar(self):
        """Seleciona PDFs para juntar"""
        arquivos, _ = QFileDialog.getOpenFileNames(self, "Selecionar PDFs para Juntar", "", "PDFs (*.pdf);;Todos os arquivos (*)")

        if arquivos:
            self.lista_pdfs.clear()
            for arquivo in arquivos:
                self.lista_pdfs.addItem(f"📁 {arquivo}")
            self.adicionar_log(f"{len(arquivos)} PDF(s) selecionado(s) para junção")

    def juntar_pdfs(self):
        """Junta os PDFs selecionados"""
        if self.lista_pdfs.count() < 2:
            QMessageBox.warning(self, "Erro", "Selecione pelo menos 2 PDFs para juntar!")
            return

        # Selecionar arquivo de saída (permite escolher local e nome)
        caminho_final, _ = QFileDialog.getSaveFileName(
            self,
            "Salvar PDF juntado como",
            os.path.join(self.pasta_saida or os.path.join(os.path.expanduser("~"), "Desktop"), "pdf_juntado.pdf"),
            "PDF (*.pdf)",
        )
        if not caminho_final:
            return

        try:
            merger = PdfMerger()

            # Adicionar cada PDF
            for i in range(self.lista_pdfs.count()):
                arquivo_pdf = self.lista_pdfs.item(i).text()
                merger.append(arquivo_pdf)

            # Salvar PDF final
            merger.write(caminho_final)
            merger.close()

            nome_arquivo_final = os.path.basename(caminho_final)
            self.adicionar_log(f"✅ PDFs juntados com sucesso: {nome_arquivo_final}")
            QMessageBox.information(self, "Sucesso", f"PDFs juntados com sucesso!\nSalvo em: {caminho_final}")

            # Limpar lista
            self.lista_pdfs.clear()

        except Exception as e:
            self.adicionar_log(f"❌ Erro ao juntar PDFs: {str(e)}")
            QMessageBox.critical(self, "Erro", f"Erro ao juntar PDFs:\n{str(e)}")

    # Drag & Drop
    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            # Aceita imagens e PDFs
            for url in event.mimeData().urls():
                arquivo = str(url.toLocalFile())
                sufixo = Path(arquivo).suffix.lower()
                if sufixo in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp", ".pdf"}:
                    event.acceptProposedAction()
                    return
        event.ignore()

    def dropEvent(self, event):
        if not event.mimeData().hasUrls():
            return

        arquivos = [str(url.toLocalFile()) for url in event.mimeData().urls()]
        imagens = [a for a in arquivos if Path(a).suffix.lower() in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"}]
        pdfs = [a for a in arquivos if Path(a).suffix.lower() == ".pdf"]

        if imagens:
            for imagem in imagens:
                self.lista_imagens.addItem(f"📁 {imagem}")
            self.adicionar_log(f"{len(imagens)} imagem(ns) adicionada(s) por drag & drop")

        if pdfs:
            # Adiciona nos PDFs para junção
            for pdf in pdfs:
                self.lista_pdfs.addItem(f"📁 {pdf}")
            self.adicionar_log(f"{len(pdfs)} PDF(s) adicionados na junção por drag & drop")

            # Se conversão especial 'De' estiver em PDF e nenhum arquivo selecionado, define-o
            try:
                if hasattr(self, "cmb_de") and "PDF" in self.cmb_de.currentText():
                    if not hasattr(self, "arquivo_especial_selecionado"):
                        self.arquivo_especial_selecionado = pdfs[0]
                        self.lbl_arquivo_selecionado.setText(f"📁 {Path(pdfs[0]).name}")
                        self.btn_converter_especial.setEnabled(True)
            except Exception:
                pass

    # === MÉTODOS PARA CONVERSÕES ESPECIAIS ===

    def atualizar_filtros_arquivo(self):
        """Atualiza os filtros de arquivo baseado na seleção 'De'"""
        tipo_de = self.cmb_de.currentText()

        # Limpar seleção anterior
        self.lbl_arquivo_selecionado.setText("Nenhum arquivo selecionado")
        self.btn_converter_especial.setEnabled(False)
        if hasattr(self, "arquivo_especial_selecionado"):
            delattr(self, "arquivo_especial_selecionado")

    def selecionar_arquivo_conversao(self):
        """Seleciona arquivo baseado no tipo escolhido no dropdown 'De'"""
        tipo_de = self.cmb_de.currentText()

        # Definir filtros baseado no tipo
        if "Excel" in tipo_de:
            filtro = "Excel Files (*.xlsx *.xls);;Todos os arquivos (*)"
            titulo = "Selecionar Arquivo Excel"
        elif "PDF" in tipo_de:
            filtro = "PDF Files (*.pdf);;Todos os arquivos (*)"
            titulo = "Selecionar Arquivo PDF"
        elif "Word" in tipo_de:
            filtro = "Word Files (*.docx);;Todos os arquivos (*)"
            titulo = "Selecionar Arquivo Word"
        elif "Imagem" in tipo_de:
            filtro = "Imagens (*.png *.jpg *.jpeg *.gif *.bmp *.tiff *.webp);;Todos os arquivos (*)"
            titulo = "Selecionar Arquivo de Imagem"
        else:
            filtro = "Todos os arquivos (*)"
            titulo = "Selecionar Arquivo"

        arquivo, _ = QFileDialog.getOpenFileName(self, titulo, "", filtro)

        if arquivo:
            self.arquivo_especial_selecionado = arquivo
            self.lbl_arquivo_selecionado.setText(f"📁 {arquivo}")
            self.btn_converter_especial.setEnabled(True)
            self.adicionar_log(f"Arquivo selecionado para conversão: {Path(arquivo).name}")

    def converter_arquivo_especial(self):
        """Converte arquivo baseado na seleção 'De' e 'Para'"""
        if not hasattr(self, "arquivo_especial_selecionado"):
            QMessageBox.warning(self, "Erro", "Selecione um arquivo primeiro!")
            return

        tipo_de = self.cmb_de.currentText()
        tipo_para = self.cmb_para.currentText()

        # Selecionar nome e pasta de saída
        tipo_de = self.cmb_de.currentText()
        tipo_para = self.cmb_para.currentText()
        sufixo = ".pdf" if "PDF" in tipo_para else (".docx" if "Word" in tipo_para else (".png" if "Imagem" in tipo_para else ".xlsx"))
        nome_padrao = Path(self.arquivo_especial_selecionado).stem + sufixo
        caminho_saida, _ = QFileDialog.getSaveFileName(
            self, "Salvar como", os.path.join(self.pasta_saida, nome_padrao), f"Arquivos (*{sufixo})"
        )
        if not caminho_saida:
            return
        pasta_saida = os.path.dirname(caminho_saida)

        try:
            if "Excel" in tipo_de and "PDF" in tipo_para:
                self._converter_excel_para_pdf(pasta_saida)
                fonte = os.path.join(pasta_saida, Path(self.arquivo_especial_selecionado).stem + ".pdf")
                if os.path.exists(fonte) and fonte != caminho_saida:
                    try:
                        os.replace(fonte, caminho_saida)
                    except Exception:
                        pass
            elif "PDF" in tipo_de and "Word" in tipo_para:
                self._converter_pdf_para_word(pasta_saida)
                fonte = os.path.join(pasta_saida, Path(self.arquivo_especial_selecionado).stem + ".docx")
                if os.path.exists(fonte) and fonte != caminho_saida:
                    try:
                        os.replace(fonte, caminho_saida)
                    except Exception:
                        pass
            elif "PDF" in tipo_de and "Imagem" in tipo_para:
                self._converter_pdf_para_imagem(pasta_saida)
                # múltiplas imagens; salva na pasta escolhida com base do nome
            elif "Imagem" in tipo_de and "PDF" in tipo_para:
                self._converter_imagem_para_pdf(pasta_saida)
                fonte = os.path.join(pasta_saida, Path(self.arquivo_especial_selecionado).stem + ".pdf")
                if os.path.exists(fonte) and fonte != caminho_saida:
                    try:
                        os.replace(fonte, caminho_saida)
                    except Exception:
                        pass
            elif "Word" in tipo_de and "PDF" in tipo_para:
                self._converter_word_para_pdf(pasta_saida)
                fonte = os.path.join(pasta_saida, Path(self.arquivo_especial_selecionado).stem + ".pdf")
                if os.path.exists(fonte) and fonte != caminho_saida:
                    try:
                        os.replace(fonte, caminho_saida)
                    except Exception:
                        pass
            elif "Excel" in tipo_de and "Word" in tipo_para:
                self._converter_excel_para_word(pasta_saida)
                fonte = os.path.join(pasta_saida, Path(self.arquivo_especial_selecionado).stem + ".docx")
                if os.path.exists(fonte) and fonte != caminho_saida:
                    try:
                        os.replace(fonte, caminho_saida)
                    except Exception:
                        pass
            else:
                QMessageBox.warning(self, "Erro", f"Conversão não suportada: {tipo_de} → {tipo_para}")
                return

        except Exception as e:
            self.adicionar_log(f"❌ Erro na conversão: {str(e)}")
            QMessageBox.critical(self, "Erro", f"Erro na conversão:\n{str(e)}")

    def _converter_excel_para_pdf(self, pasta_saida):
        """Converte Excel para PDF"""
        nome_arquivo = Path(self.arquivo_especial_selecionado).stem
        caminho_pdf = os.path.join(pasta_saida, f"{nome_arquivo}.pdf")

        workbook = openpyxl.load_workbook(self.arquivo_especial_selecionado)
        c = canvas.Canvas(caminho_pdf, pagesize=A4)
        width, height = A4

        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, height - 50, f"Planilha: {sheet_name}")

            y_position = height - 100
            for row in worksheet.iter_rows(values_only=True):
                if y_position < 100:
                    c.showPage()
                    y_position = height - 50

                x_position = 50
                for cell_value in row:
                    if cell_value is not None:
                        c.setFont("Helvetica", 10)
                        c.drawString(x_position, y_position, str(cell_value)[:50])
                        x_position += 100

                y_position -= 20

            c.showPage()

        c.save()
        self._finalizar_conversao(caminho_pdf, "Excel convertido para PDF")
        # Apagar origem se necessário
        if not self.check_manter_original_especial.isChecked():
            try:
                os.remove(self.arquivo_especial_selecionado)
            except Exception:
                pass

    def _converter_pdf_para_word(self, pasta_saida):
        """Converte PDF para Word"""
        nome_arquivo = Path(self.arquivo_especial_selecionado).stem
        caminho_word = os.path.join(pasta_saida, f"{nome_arquivo}.docx")

        images = pdf2image.convert_from_path(self.arquivo_especial_selecionado)
        doc = Document()
        doc.add_heading(f"Documento convertido de PDF: {nome_arquivo}", 0)

        for i, image in enumerate(images):
            temp_img_path = os.path.join(pasta_saida, f"temp_page_{i}.png")
            image.save(temp_img_path, "PNG")
            doc.add_heading(f"Página {i+1}", level=1)
            from docx.shared import Inches

            doc.add_picture(temp_img_path, width=Inches(6))
            os.remove(temp_img_path)

        doc.save(caminho_word)
        self._finalizar_conversao(caminho_word, "PDF convertido para Word")
        if not self.check_manter_original_especial.isChecked():
            try:
                os.remove(self.arquivo_especial_selecionado)
            except Exception:
                pass

    def _converter_pdf_para_imagem(self, pasta_saida):
        """Converte PDF para imagens"""
        nome_arquivo = Path(self.arquivo_especial_selecionado).stem
        images = pdf2image.convert_from_path(self.arquivo_especial_selecionado)

        for i, image in enumerate(images):
            caminho_imagem = os.path.join(pasta_saida, f"{nome_arquivo}_pagina_{i+1}.png")
            image.save(caminho_imagem, "PNG")

        self._finalizar_conversao(f"{len(images)} imagens", f"PDF convertido para {len(images)} imagem(ns)")
        if not self.check_manter_original_especial.isChecked():
            try:
                os.remove(self.arquivo_especial_selecionado)
            except Exception:
                pass

    def _converter_imagem_para_pdf(self, pasta_saida):
        """Converte imagem para PDF"""
        nome_arquivo = Path(self.arquivo_especial_selecionado).stem
        caminho_pdf = os.path.join(pasta_saida, f"{nome_arquivo}.pdf")

        with Image.open(self.arquivo_especial_selecionado) as img:
            if img.mode != "RGB":
                img = img.convert("RGB")
            img.save(caminho_pdf, "PDF", resolution=300.0)

        self._finalizar_conversao(caminho_pdf, "Imagem convertida para PDF")
        if not self.check_manter_original_especial.isChecked():
            try:
                os.remove(self.arquivo_especial_selecionado)
            except Exception:
                pass

    def _converter_word_para_pdf(self, pasta_saida):
        """Converte Word para PDF (simplificado)"""
        nome_arquivo = Path(self.arquivo_especial_selecionado).stem
        caminho_pdf = os.path.join(pasta_saida, f"{nome_arquivo}.pdf")

        # Para Word → PDF, vamos converter para HTML primeiro e depois para PDF
        # Esta é uma implementação simplificada
        doc = Document(self.arquivo_especial_selecionado)

        # Criar PDF básico com o texto do documento
        c = canvas.Canvas(caminho_pdf, pagesize=A4)
        width, height = A4
        y_position = height - 50

        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, y_position, f"Documento: {nome_arquivo}")
        y_position -= 30

        c.setFont("Helvetica", 12)
        for paragraph in doc.paragraphs:
            if y_position < 100:
                c.showPage()
                y_position = height - 50

            text = paragraph.text.strip()
            if text:
                # Quebrar texto em linhas se necessário
                lines = [text[i : i + 80] for i in range(0, len(text), 80)]
                for line in lines:
                    c.drawString(50, y_position, line)
                    y_position -= 15

        c.save()
        self._finalizar_conversao(caminho_pdf, "Word convertido para PDF")
        if not self.check_manter_original_especial.isChecked():
            try:
                os.remove(self.arquivo_especial_selecionado)
            except Exception:
                pass

    def _converter_excel_para_word(self, pasta_saida):
        """Converte Excel para Word"""
        nome_arquivo = Path(self.arquivo_especial_selecionado).stem
        caminho_word = os.path.join(pasta_saida, f"{nome_arquivo}.docx")

        workbook = openpyxl.load_workbook(self.arquivo_especial_selecionado)
        doc = Document()
        doc.add_heading(f"Planilha: {nome_arquivo}", 0)

        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            doc.add_heading(f"Planilha: {sheet_name}", level=1)

            # Criar tabela no Word
            table = doc.add_table(rows=1, cols=min(10, worksheet.max_column))
            table.style = "Table Grid"

            # Adicionar cabeçalhos
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(worksheet.iter_cols(max_col=10, values_only=True)):
                if i < len(hdr_cells):
                    hdr_cells[i].text = f"Coluna {i+1}"

            # Adicionar dados (limitado a 50 linhas para não sobrecarregar)
            for row_num, row in enumerate(worksheet.iter_rows(max_row=50, values_only=True)):
                if row_num > 0:  # Pular cabeçalho
                    row_cells = table.add_row().cells
                    for i, cell_value in enumerate(row):
                        if i < len(row_cells) and cell_value is not None:
                            row_cells[i].text = str(cell_value)[:50]

        doc.save(caminho_word)
        self._finalizar_conversao(caminho_word, "Excel convertido para Word")
        if not self.check_manter_original_especial.isChecked():
            try:
                os.remove(self.arquivo_especial_selecionado)
            except Exception:
                pass

    def _finalizar_conversao(self, caminho_saida, mensagem):
        """Finaliza a conversão e limpa a interface"""
        self.adicionar_log(f"✅ {mensagem}: {Path(caminho_saida).name}")
        QMessageBox.information(self, "Sucesso", f"{mensagem} com sucesso!\nSalvo em: {caminho_saida}")

        # Limpar seleção
        self.lbl_arquivo_selecionado.setText("Nenhum arquivo selecionado")
        self.btn_converter_especial.setEnabled(False)
        if hasattr(self, "arquivo_especial_selecionado"):
            delattr(self, "arquivo_especial_selecionado")

    def aplicar_tema(self, tema_escuro):
        if not PIL_AVAILABLE:
            return

        if tema_escuro:
            self.setStyleSheet(
                """
                QWidget {background-color: #2c3e50; color: white;}
                QGroupBox {font-weight: bold; border: 2px solid #7f8c8d; border-radius: 5px; margin-top: 1ex; padding-top: 10px;}
                QGroupBox::title {subcontrol-origin: margin; left: 10px; padding: 0 5px 0 5px;}
                QLabel {color: white;}
                QPushButton {border-radius: 8px; padding: 8px;}
                QComboBox {background-color: #34495e; color: white; border: 1px solid #7f8c8d; padding: 4px;}
                QLineEdit {background-color: #34495e; color: white; border: 1px solid #7f8c8d; padding: 4px;}
                QListWidget {background-color: #34495e; color: white; border: 1px solid #7f8c8d;}
                QTextEdit {background-color: #34495e; color: white; border: 1px solid #7f8c8d;}
                QCheckBox {color: white;}
            """
            )
            # Cores fixas dos botões (não mudam com tema)
            if hasattr(self, "btn_pasta_entrada"):
                self.btn_pasta_entrada.setStyleSheet("background-color: #3498db; color: white;")  # Azul - Selecionar Entrada
            if hasattr(self, "btn_pasta_saida_monitor"):
                self.btn_pasta_saida_monitor.setStyleSheet("background-color: #27ae60; color: white;")  # Verde - Selecionar Saída
            if hasattr(self, "btn_iniciar_monitor"):
                self.btn_iniciar_monitor.setStyleSheet("background-color: #2ecc71; color: white;")  # Verde - Iniciar
            if hasattr(self, "btn_pausar_monitor"):
                self.btn_pausar_monitor.setStyleSheet("background-color: #f39c12; color: white;")  # Amarelo - Pausar
            if hasattr(self, "btn_cancelar_monitor"):
                self.btn_cancelar_monitor.setStyleSheet("background-color: #e74c3c; color: white;")  # Vermelho - Cancelar
            if hasattr(self, "btn_limpar_monitor"):
                self.btn_limpar_monitor.setStyleSheet("background-color: #e74c3c; color: white;")  # Vermelho - Limpar
            if hasattr(self, "btn_selecionar_imagens"):
                self.btn_selecionar_imagens.setStyleSheet("background-color: #3498db; color: white;")  # Azul - Selecionar
            if hasattr(self, "btn_converter_imagens"):
                self.btn_converter_imagens.setStyleSheet("background-color: #2ecc71; color: white;")  # Verde - Converter
            if hasattr(self, "btn_limpar_lista"):
                self.btn_limpar_lista.setStyleSheet("background-color: #e74c3c; color: white;")  # Vermelho - Limpar
            if hasattr(self, "btn_limpar_log"):
                self.btn_limpar_log.setStyleSheet("background-color: #e74c3c; color: white;")  # Vermelho - Limpar
            if hasattr(self, "btn_pasta_saida"):
                self.btn_pasta_saida.setStyleSheet("background-color: #3498db; color: white;")  # Azul - Selecionar
            if hasattr(self, "btn_selecionar_pdfs"):
                self.btn_selecionar_pdfs.setStyleSheet("background-color: #3498db; color: white;")  # Azul - Selecionar
            if hasattr(self, "btn_juntar_pdfs"):
                self.btn_juntar_pdfs.setStyleSheet("background-color: #2ecc71; color: white;")  # Verde - Juntar
            if hasattr(self, "btn_limpar_juncao"):
                self.btn_limpar_juncao.setStyleSheet("background-color: #e74c3c; color: white;")  # Vermelho - Limpar
            # Botões das conversões especiais
            if hasattr(self, "btn_selecionar_arquivo"):
                self.btn_selecionar_arquivo.setStyleSheet("background-color: #3498db; color: white;")  # Azul - Selecionar
            if hasattr(self, "btn_converter_especial"):
                self.btn_converter_especial.setStyleSheet("background-color: #2ecc71; color: white;")  # Verde - Converter
            if hasattr(self, "btn_limpar_especial"):
                self.btn_limpar_especial.setStyleSheet("background-color: #e74c3c; color: white;")  # Vermelho - Limpar
        else:
            self.setStyleSheet(
                """
                QWidget {background-color: #ecf0f1; color: black;}
                QGroupBox {font-weight: bold; border: 2px solid #7f8c8d; border-radius: 5px; margin-top: 1ex; padding-top: 10px;}
                QGroupBox::title {subcontrol-origin: margin; left: 10px; padding: 0 5px 0 5px;}
                QLabel {color: black;}
                QPushButton {border-radius: 8px; padding: 8px;}
                QComboBox {background-color: #bdc3c7; color: black; border: 1px solid #7f8c8d; padding: 4px;}
                QLineEdit {background-color: #bdc3c7; color: black; border: 1px solid #7f8c8d; padding: 4px;}
                QListWidget {background-color: #bdc3c7; color: black; border: 1px solid #7f8c8d;}
                QTextEdit {background-color: #bdc3c7; color: black; border: 1px solid #7f8c8d;}
                QCheckBox {color: black;}
            """
            )
            # Cores fixas dos botões (não mudam com tema) - Tema Claro
            if hasattr(self, "btn_pasta_entrada"):
                self.btn_pasta_entrada.setStyleSheet("background-color: #3498db; color: white;")  # Azul - Selecionar Entrada
            if hasattr(self, "btn_pasta_saida_monitor"):
                self.btn_pasta_saida_monitor.setStyleSheet("background-color: #27ae60; color: white;")  # Verde - Selecionar Saída
            if hasattr(self, "btn_iniciar_monitor"):
                self.btn_iniciar_monitor.setStyleSheet("background-color: #2ecc71; color: white;")  # Verde - Iniciar
            if hasattr(self, "btn_pausar_monitor"):
                self.btn_pausar_monitor.setStyleSheet("background-color: #f39c12; color: white;")  # Amarelo - Pausar
            if hasattr(self, "btn_cancelar_monitor"):
                self.btn_cancelar_monitor.setStyleSheet("background-color: #e74c3c; color: white;")  # Vermelho - Cancelar
            if hasattr(self, "btn_limpar_monitor"):
                self.btn_limpar_monitor.setStyleSheet("background-color: #e74c3c; color: white;")  # Vermelho - Limpar
            if hasattr(self, "btn_selecionar_imagens"):
                self.btn_selecionar_imagens.setStyleSheet("background-color: #3498db; color: white;")  # Azul - Selecionar
            if hasattr(self, "btn_converter_imagens"):
                self.btn_converter_imagens.setStyleSheet("background-color: #2ecc71; color: white;")  # Verde - Converter
            if hasattr(self, "btn_limpar_lista"):
                self.btn_limpar_lista.setStyleSheet("background-color: #e74c3c; color: white;")  # Vermelho - Limpar
            if hasattr(self, "btn_limpar_log"):
                self.btn_limpar_log.setStyleSheet("background-color: #e74c3c; color: white;")  # Vermelho - Limpar
            if hasattr(self, "btn_pasta_saida"):
                self.btn_pasta_saida.setStyleSheet("background-color: #3498db; color: white;")  # Azul - Selecionar
            if hasattr(self, "btn_selecionar_pdfs"):
                self.btn_selecionar_pdfs.setStyleSheet("background-color: #3498db; color: white;")  # Azul - Selecionar
            if hasattr(self, "btn_juntar_pdfs"):
                self.btn_juntar_pdfs.setStyleSheet("background-color: #2ecc71; color: white;")  # Verde - Juntar
            if hasattr(self, "btn_limpar_juncao"):
                self.btn_limpar_juncao.setStyleSheet("background-color: #e74c3c; color: white;")  # Vermelho - Limpar
            # Botões das conversões especiais
            if hasattr(self, "btn_selecionar_arquivo"):
                self.btn_selecionar_arquivo.setStyleSheet("background-color: #3498db; color: white;")  # Azul - Selecionar
            if hasattr(self, "btn_converter_especial"):
                self.btn_converter_especial.setStyleSheet("background-color: #2ecc71; color: white;")  # Verde - Converter
            if hasattr(self, "btn_limpar_especial"):
                self.btn_limpar_especial.setStyleSheet("background-color: #e74c3c; color: white;")  # Vermelho - Limpar

    def __del__(self):
        """Destrutor para parar o monitoramento"""
        if hasattr(self, "watcher") and self.watcher:
            self.watcher.deleteLater()
        if hasattr(self, "timer_monitoramento") and self.timer_monitoramento:
            self.timer_monitoramento.stop()

